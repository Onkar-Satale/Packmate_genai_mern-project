"""
Smart Packing Assistant API - FastAPI Backend

This Python fastAPI service acts as the GenAI backend for the Smart Packing Assistant application. 
It receives trip details from the frontend, queries an LLM (Groq API) for a customized packing list 
in JSON format, and returns the list. It also provides functionality to save trips to a separate Node.js 
MongoDB backend and generate downloadable DOCX files.
"""

# ==========================================
# ### IMPORTS & CONFIGURATION ###
# ==========================================
import os
import json
import time
from typing import Optional
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from io import BytesIO
from fastapi import Request
from fastapi import Request
from docx import Document
import re
import requests
from groq import Groq
from pathlib import Path
from datetime import datetime, timedelta
import random
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Resolve base directory to locate the .env file containing API keys
BASE_DIR = Path(__file__).resolve().parent
ENV_PATH = BASE_DIR / ".env"

# Load environment variables from the .env file into the os.environ dictionary
load_dotenv(dotenv_path=ENV_PATH, override=True)

# Retrieve API keys from environment variables
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Ensure the Groq API key is present before starting the app.
if not GROQ_API_KEY:
    raise Exception("GROQ_API_KEY not loaded!")

GENAI_API_SECRET = os.getenv("GENAI_API_SECRET", "")

def verify_api_key(x_api_key: str = Header(None)):
    if not GENAI_API_SECRET:
        logger.warning("GENAI_API_SECRET is not set. Service is unprotected!")
    elif x_api_key != GENAI_API_SECRET:
        raise HTTPException(status_code=401, detail="Unauthorized: Invalid API Key")

# Optimize requests globally with a session
http_session = requests.Session()

# ==========================================
# ### APP INITIALIZATION & MIDDLEWARE ###
# ==========================================

# Simple in-memory cache for generated packing lists to prevent AI spam attacks
generation_cache = {}
# Cache for 10-day weather forecasts by location to prevent redundant API calls
weather_cache = {}

# Initialize the FastAPI application
app = FastAPI(title="🎒Smart Packing Assistant API")

# Rate limiting is now handled securely by the Node.js API Gateway (express-rate-limit).

# Health-check root endpoint
@app.get("/")
def root():
    """
    Root endpoint to verify the API is running successfully.
    Returns: A simple JSON message indicating live status.
    """
    return {"message": "Smart Packing Assistant API is live ✅"}

# Allow frontend origins completely defined dynamically via .env
frontend_env = os.getenv("FRONTEND_URL", "")
FRONTEND_URL = [url.strip() for url in frontend_env.split(",") if url.strip()]

# Add CORS middleware to allow the React frontend to communicate with this backend.
app.add_middleware(
    CORSMiddleware,
    allow_origins=FRONTEND_URL,
    allow_credentials=True,
    allow_methods=["*"],  # Allows all HTTP methods (GET, POST, etc.)
    allow_headers=["*"],  # Allows all headers
)

# ==========================================
# ### PYDANTIC DATA MODELS ###
# ==========================================
# These models define the expected shape of the incoming JSON requests.
# FastAPI validates incoming requests automatically based on these definitions.

class TripRequestGenerate(BaseModel):
    """
    Expected input format for generating a packing list or downloading it.
    """
    location: str = Field(..., max_length=150)
    days: int = Field(..., ge=1, le=120)
    trip_type: str = Field(..., max_length=50)
    purpose: str = Field(..., max_length=50)
    activities: str = Field(..., max_length=300)
    stay_type: str = Field(..., max_length=50)
    budget: str = Field(..., max_length=50)
    food: str = Field(..., max_length=100)
    luggage: str = Field(..., max_length=100)
    travel_type: str = Field(..., max_length=100)
    people: str = Field(..., max_length=1000)  # A flat string describing all travelers (e.g. "John, 25 years, Male")
    temperature: Optional[float] = None
    start_date: Optional[str] = Field(None, max_length=20)
    end_date: Optional[str] = Field(None, max_length=20)

class PrefetchWeatherRequest(BaseModel):
    """
    Payload for fetching weather and correcting city names
    """
    location: str = Field(..., max_length=150)


class DownloadRequest(BaseModel):
    """
    Expected input format for downloading a generated packing list.
    """
    packing_list: list

# ==========================================
# ### HELPER FUNCTIONS ###
# ==========================================

def is_section_heading(item: str) -> bool:
    item = item.strip()
    if not item:
        return False
    if item.isupper():
        return True
        
    return False

def create_docx(packing_list: list):
    """
    Creates an in-memory Microsoft Word document (.docx) containing the packing list.
    
    Args:
        packing_list (list): The list of packing items including category headers.
        
    Returns:
        BytesIO: A byte buffer containing the Word document data, ready for streaming.
    """
    buffer = BytesIO()
    doc = Document()
    
    # Add title and current date
    doc.add_heading("🎒 Smart Packing Assistant", level=1)
    
    current_date = datetime.now().strftime("%Y-%m-%d")
    doc.add_paragraph(f"Generated on: {current_date}")
            
    # Add packing list items
    for item in packing_list:
        item_stripped = item.strip()
        if not item_stripped:
            continue
        
        if is_section_heading(item_stripped):
            doc.add_heading(item_stripped, level=2)
        else:
            doc.add_paragraph(item_stripped, style="List Bullet")
        
    doc.save(buffer)
    buffer.seek(0) # Reset buffer position to the start for reading
    return buffer

def geocode_location(location: str):
    """
    Looks up latitude and longitude for a given city string using Open-Meteo's Geocoding API.
    """
    try:
        url = f"https://geocoding-api.open-meteo.com/v1/search?name={location}&count=1"
        res = http_session.get(url, timeout=5)
        if res.status_code == 200:
            data = res.json()
            if "results" in data and len(data["results"]) > 0:
                loc_data = data["results"][0]
                lat = loc_data.get("latitude")
                lon = loc_data.get("longitude")
                logger.info(f"Geocoding success for '{location}': lat={lat}, lon={lon}")
                return lat, lon
            else:
                logger.warning(f"Geocoding found no results for '{location}'")
                return None, None
        else:
            logger.error(f"Geocoding API returned status {res.status_code}")
            return None, None
    except Exception as e:
        logger.error(f"Geocoding error: {e}")
        return None, None

def prefetch_and_cache_weather(location: str):
    """
    Fetches the 16-day weather forecast using Open-Meteo API and caches it.
    Returns the current/average temperature for the provided destination for response compatibility.
    """
    loc_key = location.lower().strip()
    lat, lon = geocode_location(loc_key)
    if lat is None or lon is None:
        logger.warning(f"Skipping Open-Meteo forecast fetch due to geocoding failure for '{location}'")
        return None

    try:
        url = f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&daily=temperature_2m_max,temperature_2m_min&forecast_days=16&timezone=auto"
        res = http_session.get(url, timeout=5)
        
        if res.status_code == 200:
            data = res.json()
            api_temps = {}
            current_temp = None
            last_api_temp = None
            
            if "daily" in data:
                daily = data["daily"]
                time_list = daily.get("time", [])
                max_temps = daily.get("temperature_2m_max", [])
                min_temps = daily.get("temperature_2m_min", [])
                
                for i, dt_key in enumerate(time_list):
                    t_max = max_temps[i]
                    t_min = min_temps[i]
                    if t_max is not None and t_min is not None:
                        avg_temp = round((t_max + t_min) / 2.0, 1)
                        api_temps[dt_key] = avg_temp
                        last_api_temp = avg_temp
                        if i == 0:
                            current_temp = avg_temp
                            
            if current_temp is None and last_api_temp is not None:
                current_temp = last_api_temp
                
            weather_cache[loc_key] = {
                "api_temps": api_temps,
                "last_api_temp": last_api_temp,
                "fetched_at": time.time()
            }
            
            logger.info(f"Weather prefetched from Open-Meteo for '{location}' and cached successfully.")
            
            return current_temp
        else:
            logger.error(f"Open-Meteo API /forecast returned status {res.status_code}")
            return None
    except Exception as e:
        logger.error(f"Open-Meteo API error: {e}")
        return None

def compute_full_trip_weather(data: dict) -> str:
    location = data.get("location", "").lower().strip()
    start_date_str = data.get("start_date")
    end_date_str = data.get("end_date")
    trip_days = int(data.get("days", 1))
    
    fallback_temp = data.get('temperature')
    if fallback_temp is None:
        fallback_temp = random.randint(20, 30)
        logger.info(f"No temp provided, using random fallback: {fallback_temp}°C")

    try:
        # Fallback if dates are missing: use today + trip_days
        if not start_date_str or not end_date_str:
            logger.info(f"Dates missing for {location}, falling back to today + {trip_days} days")
            start_date = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
            total_days = trip_days
        else:
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d")
            total_days = (end_date - start_date).days + 1
            if total_days <= 0: total_days = 1
        
        date_list = [start_date + timedelta(days=i) for i in range(total_days)]
        
        cached_weather = weather_cache.get(location)
        if not cached_weather:
            # Graceful fetch in case prefetch wasn't called or cache was lost
            logger.info("Weather cache miss in compute_full_trip_weather, fetching now.")
            _ = prefetch_and_cache_weather(location)
            cached_weather = weather_cache.get(location)
            
        if cached_weather and cached_weather.get("api_temps"):
            api_temps = cached_weather["api_temps"]
            # Retrieve 16th day (last available) temperature to act as our base
            sorted_dates = sorted(api_temps.keys())
            sixteenth_day_temp = api_temps[sorted_dates[-1]] if sorted_dates else fallback_temp
        else:
            api_temps = {}
            sixteenth_day_temp = fallback_temp

        forecast_api_lines = []
        forecast_extra_lines = []
        current_temp = sixteenth_day_temp

        for dt in date_list:
            dt_str = dt.strftime("%Y-%m-%d")
            
            if dt_str in api_temps:
                current_temp = api_temps[dt_str]
                forecast_api_lines.append(f"{dt_str} → {current_temp}°C (Prefetched)")
            else:
                drift = random.choices([-2, -1, 0, 1, 2], weights=[0.1, 0.35, 0.1, 0.35, 0.1])[0]
                new_temp = current_temp + drift

                if new_temp > sixteenth_day_temp + 5:
                    new_temp = sixteenth_day_temp + 5
                elif new_temp < sixteenth_day_temp - 5:
                    new_temp = sixteenth_day_temp - 5

                current_temp = round(new_temp, 1)
                forecast_extra_lines.append(f"{dt_str} → {current_temp}°C (Generated Drift)")
        
        all_lines = forecast_api_lines + forecast_extra_lines
        
        formatted_output = f"\n================ FULL DAY-WISE TEMPERATURE MAPPING FOR '{location.upper()}' ================\n"
        formatted_output += "\n".join(all_lines)
        formatted_output += "\n=================================================================================="
        
        logger.info(formatted_output)
        
        clean_lines = [line.replace(" (Prefetched)", "").replace(" (Generated Drift)", "") for line in all_lines]
        weather_text = "Day-wise temperature forecast:\n" + "\n".join(clean_lines)

        return weather_text

    except Exception as e:
        logger.error(f"Error generating full trip weather: {e}")
        return f"Average temperature: {fallback_temp}°C"

def generate_packing_data(data: dict):
    """
    The core logic integrating with the Groq Large Language Model.
    Sends environmental and traveler data to generate a complete packing list.
    
    Args:
        data (dict): The dictionary representation of the TripRequestGenerate model.
        
    Returns:
        dict: A parsed JSON response matching the required structure:
            { "packing_list": ["header", "item", ...] }
    """
    # Initialize the Groq client to call the LLM
    client = Groq(api_key=GROQ_API_KEY)

    # Gather dynamic day-wise temperature array from precomputed exact API logic
    temp_info = compute_full_trip_weather(data)
    
    # Define the system prompt guiding the AI's behavior, establishing rules, and restricting the output format
    system_prompt = """
========================
SMART PACKING ASSISTANT - DETAILED AI-GENERATED PACKING LIST
========================

You are a senior professional travel planner and packing consultant. Your task is to generate a complete, structured, and practical packing list for any traveler or group based entirely on the provided trip and traveler information. 
⚠️ IMPORTANT:
- Always include **all 12 mandatory sections** listed below in the exact order, even if minimal items are needed.
- Quantities, emojis, and items must be **dynamically decided** based on input data.
- Use **standardized emojis per section** consistently.
- Consider **traveler-specific details**: age, gender, medical notes, dietary restrictions, chronic conditions.
- Consider **trip details**: location, activities, duration, accommodation, budget, luggage style.
- Include optional, backup, and emergency items for all travelers.
- EXTREMELY IMPORTANT: Take the provided day-by-day temperatures VERY seriously. Evaluate the exact high and low bounds and precisely adjust clothing (thermal wear, summer wear, thin layers) to perfectly match the temperature swings. Do NOT recommend winter items for hot days, or summer items for freezing days.
- Ensure **practicality**: only include items travelers can realistically carry.
- Output **only the packing list in a flat JSON list**, do not include greetings, explanations, or nested arrays.


========================
MANDATORY SECTIONS (ALL CAPS)
========================
1. DOCUMENTS
2. CLOTHING
3. FOOTWEAR
4. TOILETRIES & PERSONAL CARE
5. ELECTRONICS & GADGETS
6. MEDICAL & HEALTH
7. ACCESSORIES
8. FOOD & SNACKS
9. ACTIVITY-SPECIFIC ITEMS
10. MISCELLANEOUS
11. WEATHER-SPECIFIC ITEMS
12. SAFETY & EMERGENCY ITEMS

========================
GENERAL INSTRUCTIONS
========================
- Include **quantities appropriate** for number of days, travelers, and laundry availability.
- Include **all traveler-specific adjustments**: age, gender, medical notes, dietary restrictions.
- Include **weather-appropriate items** based on destination temperature and season.
- Include **electronics and accessories** according to traveler details (phones, laptops, cameras, chargers).
- Include **food, snacks, and hydration** as per traveler type (solo, kids, elders, family).
- Include **backup and emergency items**: extra socks, ch+argers, first-aid, medications.
- For multi-person trips, include items for **each traveler individually**.
- **Do not use placeholders** like "stuff" or "if needed".
- Always use **consistent emojis per section**.
- Output must be a **flat JSON list**, not nested arrays.
- Ensure all items are **practical, ready-to-go, and realistic**.
- **Every section MUST have at least one item.** If no specific activities are provided, add general exploration items (e.g. "Daypack 🎒", "Comfortable walking wear 👕") to ACTIVITY-SPECIFIC ITEMS so it is never empty.

========================
OUTPUT FORMAT
========================
Respond ONLY with **JSON object** in this exact format:

{
  "packing_list": [
    "DOCUMENTS",
    "Passport 📄 (1 per traveler)",
    "Visa 🛂 (if required)",
    "Flight tickets 🎟️",
    "Hotel booking confirmation 🏨",
    "Travel insurance card 🏥",
    ...
    "CLOTHING",
    "T-shirts 👕 (N)",
    "Jeans / Pants 👖 (N)",
    "Shorts 🩳 (N)",
    "Sleepwear 😴 (N)",
    ...
    "FOOTWEAR",
    "Walking shoes 👟",
    "Sandals 🩴",
    ...
    "TOILETRIES & PERSONAL CARE",
    "Toothbrush 🪥",
    "Toothpaste 🪥",
    "Soap / Body wash 🧼",
    ...
    "MEDICAL & HEALTH",
    "First aid kit 🩹",
    "Prescription medication 💊",
    "Fever medication 💊",
    ...
    "ACCESSORIES",
    "Sunglasses 🕶️",
    "Backpack 🎒",
    "Umbrella ☔️",
    ...
    "ELECTRONICS & GADGETS",
    "Smartphone & charger 📱",
    "Laptop & charger 💻",
    "Camera 📷",
    ...
    "FOOD & SNACKS",
    "Non-perishable snacks 🍫",
    "Energy bars 🍿",
    "Reusable water bottle 💧",
    ...
    "ACTIVITY-SPECIFIC ITEMS",
    "Gym clothes 🏋️",
    "Swimwear 🏊",
    "Trekking shoes 🥾",
    ...
    "MISCELLANEOUS",
    "Notebook 📝",
    "Guidebook 📘",
    ...
    "WEATHER-SPECIFIC ITEMS",
    "Thermal wear 🧥 (cold weather)",
    "Sunscreen 🌞 (hot weather)",
    "Raincoat ☔️ (rainy weather)",
    ...
    "SAFETY & EMERGENCY ITEMS",
    "Travel locks 🔒",
    "Power bank 🔋",
    "Flashlight 🔦",
    "Emergency contact info 📇"
  ]
}
- Do not use nested arrays. Use a **flat list**.
- Each **section header is ALL CAPS, max 30 characters, no emojis**.
- Each **item under a section must include emoji** and quantity if relevant.
- Generate **all items dynamically**. Do not use placeholders like "if needed" or "stuff".
"""

    # Populate the dynamic data for the current request
    user_prompt = f"""
Location: {data['location']}
Duration: {data['days']} days
Trip Type: {data['trip_type']}
Purpose: {data['purpose']}
Activities: {data['activities']}
Stay Type: {data['stay_type']}
Budget: {data['budget']}
Food Preference: {data['food']}
Luggage Style: {data['luggage']}
Travel Mode: {data['travel_type']}
Travelers details: {data['people']}


Include all packing items dynamically based on:
- Traveler-specific medical notes, age, gender
- Weather conditions and temperature
- Trip duration, activities, and accommodation
- Budget and luggage style
- Ensure all 12 mandatory sections are included
- Use consistent emojis and realistic quantities

Environmental Context:
{temp_info}. Use this to determine weather-appropriate items.

Generate the JSON packing list now.
"""

    # Make the call to the Groq API model
    # We specify response_format={"type": "json_object"} to strictly enforce a JSON response structure.
    res = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0.4,
        max_completion_tokens=2000,
        response_format={"type": "json_object"}
    )
    
    try:
        content = res.choices[0].message.content
        
        try:
            result = json.loads(content)
        except json.JSONDecodeError as exc:
            logger.error(f"JSON decode error from Groq: {exc}\nRaw Content: {content[:200]}")
            raise HTTPException(status_code=500, detail="Invalid JSON response from AI")
        
        # Ensure packing_list items are clean from stray bullet points or numbering if the AI still added them
        clean_list = []
        for line in result.get("packing_list", []):
            if not isinstance(line, str):
                continue
            # Regex to remove bullets or numbering at the start of a response item
            clean_line = re.sub(r"^[•\-*\d.]+\s*", "", line).strip()
            if clean_line:
                clean_list.append(clean_line)
                
        logger.info(f"Successfully generated packing list containing {len(clean_list)} items")
        return {"packing_list": clean_list}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed processing AI response: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate packing list: {str(e)}")

# ==========================================
# ### API ENDPOINTS ###
# ==========================================

@app.post("/prefetch-weather", dependencies=[Depends(verify_api_key)])
def api_prefetch_weather(req: PrefetchWeatherRequest):
    """
    Endpoint to correct the city name using Groq, then prefetch the temperature.
    """
    client = Groq(api_key=GROQ_API_KEY)
    
    # 1. Correct city with Groq (Fastest model for simple NLP extraction)
    prompt = f"""
You are a strict location corrector.

Task:
Fix the spelling of the given city or country name.

Input:
{req.location}

Rules:
1. Return ONLY the corrected name.
2. Do NOT add any extra words, punctuation, or explanation.
3. Output must be a single valid city or country name.
4. If input is already correct, return it unchanged.
5. If the input is invalid or cannot be corrected, return exactly: INVALID
""" 
    try:
        res = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=20
        )
        raw_response = res.choices[0].message.content.strip()
        
        # Fallback if Groq ignored instructions and wrote a conversational apology
        if len(raw_response.split()) > 4:
            corrected_city = req.location
        else:
            corrected_city = raw_response.strip(".,'\"")
            
    except Exception as e:
        logger.error(f"Groq city correction error: {e}")
        corrected_city = req.location # Fallback to original
        
    # 2. Fetch API forecast and populate cache for the corrected city
    temp = prefetch_and_cache_weather(corrected_city)
    
    return {"original": req.location, "location": corrected_city, "temperature": temp}


@app.post("/generate-packing-list", dependencies=[Depends(verify_api_key)])
# Node.js API gateway handles rate limiting
def api_generate_packing_list(request: Request, trip: TripRequestGenerate):
    """
    Primary API Endpoint to generate an AI-driven packing list.
    
    Expected Body (TripRequestGenerate format): 
      JSON with location, days, trip_type, budget, food, luggage, people, etc.
      
    Returns: 
      JSON response containing:
      {
         "packing_list": ["SECTION", "Item", "Item", ...]
      }
    """
    data = trip.dict()
    
    cache_key = json.dumps(data, sort_keys=True)
    current_time = time.time()
    
    # Check cache first to save AI calls
    if cache_key in generation_cache:
        cached_data, timestamp = generation_cache[cache_key]
        if current_time - timestamp < 300: # 5 minute TTL
            logger.info("Serving generated list from cache.")
            return cached_data
        else:
            del generation_cache[cache_key]

    ai_result = generate_packing_data(data)
    generation_cache[cache_key] = (ai_result, current_time)
    
    return ai_result

@app.post("/download-packing-list", dependencies=[Depends(verify_api_key)])
def api_download_packing_list(req: DownloadRequest):
    """
    Endpoint that packages the provided packing list directly into a downloadable .docx file.
    
    Expected Body: 
      JSON with packing_list.
      
    Returns:
      StreamingResponse serving a Word Document file as an attachment.
    """
    # Create DOCX and stream it back
    doc = create_docx(req.packing_list)
    return StreamingResponse(
        doc,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Smart_Packing_List.docx"}
    )


# ==========================================
# ### APPLICATION ENTRY POINT ###
# ==========================================
if __name__ == "__main__":
    import uvicorn
    # Start the server using Uvicorn. Enables hot-reload during active development.
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=int(os.environ.get("PORT", 5001)), # Changed to 5001 to prevent Node port collision
        reload=False
    )
